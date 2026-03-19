#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import re
import sqlite3
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

try:
    from ops import workspace_hub_project
except ImportError:  # pragma: no cover
    import workspace_hub_project  # type: ignore


EXCLUDED_DIRS = {
    ".git",
    ".obsidian",
    "node_modules",
    "dist",
    "build",
    ".next",
    "coverage",
    "target",
    ".venv",
    "venv",
    "__pycache__",
    "runtime",
    "logs",
}
EXCLUDED_FILE_PREFIXES = ("~$", ".~")
TEXT_EXTS = {".md", ".txt", ".rst", ".adoc", ".csv"}
RICH_EXTS = {".pdf", ".docx", ".xlsx"}
SUPPORTED_EXTS = TEXT_EXTS | RICH_EXTS
MAX_CONTENT_BYTES = 2 * 1024 * 1024
REGISTRY_RE = re.compile(
    r"<!-- PROJECT_REGISTRY_DATA_START -->\s*```json\s*(.*?)\s*```\s*<!-- PROJECT_REGISTRY_DATA_END -->",
    re.S,
)
DOC_TYPE_WEIGHT = {
    "topic-board": 0,
    "project-board": 1,
    "system-doc": 2,
    "global-board": 3,
    "project-summary": 4,
    "report": 5,
    "daily-log": 6,
    "project-doc": 7,
    "other": 8,
}
SOURCE_GROUP_WEIGHT = {
    "truth": 0,
    "system-doc": 1,
    "report": 2,
    "deliverable": 3,
    "project-doc": 4,
    "daily-log": 5,
    "other": 6,
}
DELIVERABLE_PATH_MARKERS = ("deliverable", "deliverables", "交付")
MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")


def workspace_root() -> Path:
    return Path(os.environ.get("WORKSPACE_HUB_ROOT", str(workspace_hub_project.DEFAULT_WORKSPACE_ROOT)))


def vault_root() -> Path:
    return Path(
        os.environ.get(
            "WORKSPACE_HUB_VAULT_ROOT",
            str(workspace_hub_project.DEFAULT_LOCAL_VAULT_ROOT),
        )
    )


def reports_root() -> Path:
    return Path(os.environ.get("WORKSPACE_HUB_REPORTS_ROOT", str(workspace_root() / "reports")))


def projects_root() -> Path:
    return Path(os.environ.get("WORKSPACE_HUB_PROJECTS_ROOT", str(workspace_root() / "projects")))


def runtime_root() -> Path:
    explicit = os.environ.get("WORKSPACE_HUB_RUNTIME_ROOT", "").strip()
    if explicit:
        return Path(explicit)
    current_root = workspace_root()
    current_runtime = current_root / "runtime"
    worktrees_root = current_root.parent
    if worktrees_root.name == "workspace-hub-worktrees":
        canonical_runtime = workspace_hub_project.DEFAULT_WORKSPACE_ROOT / "runtime"
        if canonical_runtime.exists():
            return canonical_runtime
    return current_runtime


def retrieval_root() -> Path:
    return runtime_root() / "retrieval"


def index_db_path() -> Path:
    return retrieval_root() / "index.sqlite"


def state_path() -> Path:
    return retrieval_root() / "state.json"


def iso_now() -> str:
    import datetime as dt

    return dt.datetime.now(dt.timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def load_state() -> dict[str, Any]:
    try:
        return json.loads(state_path().read_text(encoding="utf-8"))
    except FileNotFoundError:
        return {"last_build_at": None, "last_sync_at": None, "doc_count": 0, "dirty_count": 0}
    except json.JSONDecodeError:
        return {"last_build_at": None, "last_sync_at": None, "doc_count": 0, "dirty_count": 0}


def save_state(data: dict[str, Any]) -> None:
    retrieval_root().mkdir(parents=True, exist_ok=True)
    write_text(state_path(), json.dumps(data, ensure_ascii=False, indent=2) + "\n")


def read_registry_project_names() -> list[str]:
    registry = vault_root() / "PROJECT_REGISTRY.md"
    try:
        text = registry.read_text(encoding="utf-8")
    except FileNotFoundError:
        return []
    match = REGISTRY_RE.search(text)
    if not match:
        return []
    try:
        data = json.loads(match.group(1))
    except json.JSONDecodeError:
        return []
    return sorted({str(item.get("project_name", "")).strip() for item in data if item.get("project_name")}, key=len, reverse=True)


def known_project_names() -> list[str]:
    names = set(read_registry_project_names())
    root = projects_root()
    if root.exists():
        for child in root.iterdir():
            if child.is_dir():
                names.add(child.name)
    return sorted((name for name in names if name), key=len, reverse=True)


def match_project_prefix(name: str) -> str:
    for project_name in known_project_names():
        if name == project_name or name.startswith(project_name + "-"):
            return project_name
    return ""


def iter_candidate_files() -> Iterable[Path]:
    roots = [vault_root(), reports_root(), projects_root()]
    for root in roots:
        if not root.exists():
            continue
        for path in root.rglob("*"):
            if not path.is_file():
                continue
            if any(part in EXCLUDED_DIRS for part in path.parts):
                continue
            if path.name.startswith(EXCLUDED_FILE_PREFIXES):
                continue
            if path.suffix.lower() not in SUPPORTED_EXTS:
                continue
            yield path


def file_hash(path: Path) -> str:
    h = hashlib.sha1()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def truncate_text(text: str, *, budget_bytes: int = MAX_CONTENT_BYTES) -> tuple[str, str]:
    encoded = text.encode("utf-8")
    if len(encoded) <= budget_bytes:
        return text, "ok" if text.strip() else "no_text"
    clipped = encoded[:budget_bytes]
    while True:
        try:
            result = clipped.decode("utf-8")
            break
        except UnicodeDecodeError:
            clipped = clipped[:-1]
    return result, "truncated"


def extract_text_file(path: Path) -> tuple[str, str]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return truncate_text(text)


def extract_pdf(path: Path) -> tuple[str, str]:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text:
            parts.append(page_text)
    text = "\n".join(parts).strip()
    return truncate_text(text)


def extract_docx(path: Path) -> tuple[str, str]:
    document = Document(str(path))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return truncate_text("\n".join(parts))


def extract_xlsx(path: Path) -> tuple[str, str]:
    workbook = load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"# sheet:{sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(item).strip() for item in row if item is not None and str(item).strip()]
            if values:
                parts.append(" | ".join(values))
    return truncate_text("\n".join(parts))


def extract_csv(path: Path) -> tuple[str, str]:
    parts: list[str] = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
        reader = csv.reader(handle)
        for row in reader:
            values = [item.strip() for item in row if item.strip()]
            if values:
                parts.append(" | ".join(values))
    return truncate_text("\n".join(parts))


def extract_text(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt", ".rst", ".adoc"}:
        return extract_text_file(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".xlsx":
        return extract_xlsx(path)
    if suffix == ".csv":
        return extract_csv(path)
    return "", "unsupported"


def derive_title(path: Path, text: str) -> str:
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            return stripped.lstrip("#").strip()
        if stripped:
            return stripped[:120]
    return path.stem


def derive_heading_and_lines(path: Path, text: str, title: str) -> tuple[str, int, int]:
    lines = text.splitlines()
    if not lines:
        return title, 0, 0

    first_nonempty = 0
    heading = title
    for index, raw_line in enumerate(lines, start=1):
        stripped = raw_line.strip()
        if not stripped:
            continue
        if not first_nonempty:
            first_nonempty = index
        if stripped.startswith("#"):
            heading = stripped.lstrip("#").strip() or title
            return heading, index, len(lines)
        if heading == title:
            heading = stripped[:120]
    return heading or title, first_nonempty or 1, len(lines)


def parse_query_keywords(query: str) -> list[str]:
    keywords: list[str] = []
    seen: set[str] = set()
    query = query.strip().lower()
    if not query:
        return []
    for run in re.findall(r"[\u4e00-\u9fff]+|[a-z0-9_./-]+", query):
        if re.fullmatch(r"[\u4e00-\u9fff]+", run):
            for size in (2, 1):
                if len(run) < size:
                    continue
                for index in range(len(run) - size + 1):
                    token = run[index : index + size]
                    if token not in seen:
                        seen.add(token)
                        keywords.append(token)
        else:
            token = run.strip()
            if len(token) >= 2 and token not in seen:
                seen.add(token)
                keywords.append(token)
    return keywords


def chunk_text_sections(
    text: str,
    title: str,
    *,
    chunk_size: int = 1200,
    overlap: int = 150,
) -> list[dict[str, Any]]:
    lines = text.splitlines()
    if not lines:
        return [{"heading": title, "text": "", "line_start": 0, "line_end": 0}]

    sections: list[dict[str, Any]] = []
    current_heading = title
    current_lines: list[str] = []
    section_start = 1

    for index, line in enumerate(lines, start=1):
        heading_match = MARKDOWN_HEADING_RE.match(line)
        if heading_match:
            if current_lines:
                sections.append(
                    {
                        "heading": current_heading or title,
                        "lines": current_lines,
                        "line_start": section_start,
                    }
                )
            current_heading = heading_match.group(2).strip() or title
            current_lines = [line]
            section_start = index
            continue
        if not current_lines:
            section_start = index
        current_lines.append(line)

    if current_lines:
        sections.append(
            {
                "heading": current_heading or title,
                "lines": current_lines,
                "line_start": section_start,
            }
        )

    chunks: list[dict[str, Any]] = []
    for section in sections:
        section_text = "\n".join(section["lines"])
        if len(section_text) <= chunk_size:
            chunks.append(
                {
                    "heading": section["heading"],
                    "text": section_text,
                    "line_start": section["line_start"],
                    "line_end": section["line_start"] + len(section["lines"]) - 1,
                }
            )
            continue
        char_pos = 0
        line_index = 0
        while char_pos < len(section_text):
            end_pos = min(char_pos + chunk_size, len(section_text))
            chunk_text = section_text[char_pos:end_pos]
            chunk_lines = chunk_text.splitlines() or [chunk_text]
            line_start = section["line_start"] + line_index
            line_end = line_start + len(chunk_lines) - 1
            chunks.append(
                {
                    "heading": section["heading"],
                    "text": chunk_text,
                    "line_start": line_start,
                    "line_end": line_end,
                }
            )
            if end_pos >= len(section_text):
                break
            advance = chunk_size - overlap
            advanced_text = section_text[char_pos : char_pos + advance]
            line_index += max(0, advanced_text.count("\n"))
            char_pos += advance

    return chunks


def score_section_chunk(chunk: dict[str, Any], keywords: list[str]) -> int:
    text_lower = str(chunk.get("text", "")).lower()
    heading_lower = str(chunk.get("heading", "")).lower()
    score = 0
    for keyword in keywords:
        if not keyword:
            continue
        start = 0
        while True:
            start = text_lower.find(keyword, start)
            if start == -1:
                break
            score += 1
            start += len(keyword)
        start = 0
        while True:
            start = heading_lower.find(keyword, start)
            if start == -1:
                break
            score += 2
            start += len(keyword)
    return score


def build_chunk_excerpt(chunk_text: str, keywords: list[str], *, budget: int = 320) -> str:
    normalized = re.sub(r"\s+", " ", chunk_text).strip()
    if not normalized:
        return ""
    ordered_keywords = [keyword for keyword in sorted(keywords, key=len, reverse=True) if keyword]
    hit_at = -1
    hit_len = 0
    lower_text = normalized.lower()
    for keyword in ordered_keywords:
        position = lower_text.find(keyword)
        if position != -1:
            hit_at = position
            hit_len = len(keyword)
            break
    if hit_at == -1:
        return normalized[:budget]
    start = max(0, hit_at - 80)
    end = min(len(normalized), hit_at + max(hit_len, 1) + 200)
    excerpt = normalized[start:end]
    if start > 0:
        excerpt = "…" + excerpt
    if end < len(normalized):
        excerpt = excerpt + "…"
    return excerpt[: budget + 2]


def select_best_chunk(
    *,
    content: str,
    title: str,
    fallback_heading: str,
    fallback_line_start: int,
    fallback_line_end: int,
    base_excerpt: str,
    query: str,
) -> tuple[str, int, int, str]:
    keywords = parse_query_keywords(query)
    if not content.strip() or not keywords:
        return fallback_heading, fallback_line_start, fallback_line_end, base_excerpt
    chunks = chunk_text_sections(content, title)
    best_chunk: dict[str, Any] | None = None
    best_score = 0
    for chunk in chunks:
        score = score_section_chunk(chunk, keywords)
        if score > best_score:
            best_score = score
            best_chunk = chunk
    if not best_chunk or best_score <= 0:
        return fallback_heading, fallback_line_start, fallback_line_end, base_excerpt
    return (
        str(best_chunk.get("heading") or fallback_heading or title),
        int(best_chunk.get("line_start") or fallback_line_start or 0),
        int(best_chunk.get("line_end") or fallback_line_end or 0),
        build_chunk_excerpt(str(best_chunk.get("text", "")), keywords) or base_excerpt,
    )


def classify_source_group(doc_type: str, path: Path | str) -> str:
    resolved = Path(path)
    lower_parts = [part.lower() for part in resolved.parts]
    if doc_type in {"topic-board", "project-board", "global-board"}:
        return "truth"
    if doc_type == "system-doc":
        return "system-doc"
    if doc_type == "report":
        return "report"
    if any(marker in part for part in lower_parts for marker in DELIVERABLE_PATH_MARKERS):
        return "deliverable"
    if doc_type in {"project-summary", "project-doc"}:
        return "project-doc"
    if doc_type == "daily-log":
        return "daily-log"
    return "other"


def is_within(path: str | Path, root: str | Path) -> bool:
    child_path = Path(path).resolve(strict=False)
    parent_path = Path(root).resolve(strict=False)
    try:
        child_path.relative_to(parent_path)
        return True
    except ValueError:
        return False


def is_hotset_path(path: str | Path, hotset_paths: Iterable[str] | None) -> bool:
    if not hotset_paths:
        return False
    target = str(Path(path).resolve(strict=False))
    return any(is_within(target, hotset_path) for hotset_path in hotset_paths if hotset_path)


def classify_document(path: Path) -> dict[str, str]:
    path = path.resolve()
    vault = vault_root().resolve()
    reports = reports_root().resolve()
    projects = projects_root().resolve()
    result = {"doc_type": "other", "project_name": "", "topic_name": ""}
    if vault in path.parents or path == vault:
        rel = path.relative_to(vault)
        if rel.parts[0] == "01_working":
            name = path.stem
            if name.endswith("-项目板"):
                result["doc_type"] = "project-board"
                result["project_name"] = name[: -len("-项目板")]
            elif name.endswith("-跟进板"):
                result["doc_type"] = "topic-board"
                base = name[: -len("-跟进板")]
                project_name = match_project_prefix(base)
                result["project_name"] = project_name
                if project_name:
                    topic = base[len(project_name) :].lstrip("-")
                    result["topic_name"] = topic
            else:
                result["doc_type"] = "project-doc"
                result["project_name"] = match_project_prefix(name)
        elif rel.parts[:2] == ("03_semantic", "systems"):
            result["doc_type"] = "system-doc"
        elif rel.parts[:2] == ("03_semantic", "projects"):
            if path.name == "README.md":
                result["doc_type"] = "system-doc"
            else:
                stem = path.stem
                project_name = match_project_prefix(stem) or stem
                if stem == project_name:
                    result["doc_type"] = "project-summary"
                else:
                    result["doc_type"] = "project-doc"
                result["project_name"] = project_name
        elif rel.parts[:2] == ("02_episodic", "daily"):
            result["doc_type"] = "daily-log"
        elif path.name in {"PROJECT_REGISTRY.md", "ACTIVE_PROJECTS.md", "NEXT_ACTIONS.md"} or rel.parts[0] == "07_dashboards":
            result["doc_type"] = "global-board"
        else:
            result["doc_type"] = "other"
    elif reports in path.parents or path == reports:
        result["doc_type"] = "report"
    elif projects in path.parents or path == projects:
        rel = path.relative_to(projects)
        if rel.parts:
            result["project_name"] = rel.parts[0]
        result["doc_type"] = "project-doc"
    return result


def ensure_db(conn: sqlite3.Connection) -> None:
    conn.executescript(
        """
        CREATE TABLE IF NOT EXISTS docs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            path TEXT NOT NULL UNIQUE,
            title TEXT NOT NULL,
            doc_type TEXT NOT NULL,
            source_group TEXT NOT NULL DEFAULT 'other',
            heading TEXT NOT NULL DEFAULT '',
            line_start INTEGER NOT NULL DEFAULT 0,
            line_end INTEGER NOT NULL DEFAULT 0,
            project_name TEXT NOT NULL,
            topic_name TEXT NOT NULL,
            mtime REAL NOT NULL,
            size INTEGER NOT NULL,
            content_hash TEXT NOT NULL,
            extract_status TEXT NOT NULL,
            content TEXT NOT NULL,
            indexed_at TEXT NOT NULL
        );
        CREATE VIRTUAL TABLE IF NOT EXISTS docs_fts USING fts5(
            title,
            content,
            path UNINDEXED,
            doc_type UNINDEXED,
            project_name UNINDEXED,
            topic_name UNINDEXED,
            tokenize='unicode61'
        );
        """
    )
    columns = {str(row["name"]) for row in conn.execute("PRAGMA table_info(docs)").fetchall()}
    if "source_group" not in columns:
        conn.execute("ALTER TABLE docs ADD COLUMN source_group TEXT NOT NULL DEFAULT 'other'")
    if "heading" not in columns:
        conn.execute("ALTER TABLE docs ADD COLUMN heading TEXT NOT NULL DEFAULT ''")
    if "line_start" not in columns:
        conn.execute("ALTER TABLE docs ADD COLUMN line_start INTEGER NOT NULL DEFAULT 0")
    if "line_end" not in columns:
        conn.execute("ALTER TABLE docs ADD COLUMN line_end INTEGER NOT NULL DEFAULT 0")


def db_connect() -> sqlite3.Connection:
    retrieval_root().mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(index_db_path())
    conn.row_factory = sqlite3.Row
    ensure_db(conn)
    return conn


def upsert_document(conn: sqlite3.Connection, path: Path) -> dict[str, Any]:
    metadata = classify_document(path)
    try:
        text, status = extract_text(path)
    except Exception:
        text, status = "", "error"
    title = derive_title(path, text)
    heading, line_start, line_end = derive_heading_and_lines(path, text, title)
    source_group = classify_source_group(metadata["doc_type"], path)
    stat = path.stat()
    content_hash = file_hash(path)
    now = iso_now()
    existing = conn.execute("SELECT id FROM docs WHERE path = ?", (str(path),)).fetchone()
    if existing:
        doc_id = int(existing["id"])
        conn.execute(
            """
            UPDATE docs
            SET title = ?, doc_type = ?, source_group = ?, heading = ?, line_start = ?, line_end = ?, project_name = ?, topic_name = ?, mtime = ?, size = ?, content_hash = ?, extract_status = ?, content = ?, indexed_at = ?
            WHERE id = ?
            """,
            (
                title,
                metadata["doc_type"],
                source_group,
                heading,
                line_start,
                line_end,
                metadata["project_name"],
                metadata["topic_name"],
                stat.st_mtime,
                stat.st_size,
                content_hash,
                status,
                text,
                now,
                doc_id,
            ),
        )
    else:
        cursor = conn.execute(
            """
            INSERT INTO docs(
                path,
                title,
                doc_type,
                source_group,
                heading,
                line_start,
                line_end,
                project_name,
                topic_name,
                mtime,
                size,
                content_hash,
                extract_status,
                content,
                indexed_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                str(path),
                title,
                metadata["doc_type"],
                source_group,
                heading,
                line_start,
                line_end,
                metadata["project_name"],
                metadata["topic_name"],
                stat.st_mtime,
                stat.st_size,
                content_hash,
                status,
                text,
                now,
            ),
        )
        doc_id = int(cursor.lastrowid)
    conn.execute("DELETE FROM docs_fts WHERE rowid = ?", (doc_id,))
    conn.execute(
        """
        INSERT INTO docs_fts(rowid, title, content, path, doc_type, project_name, topic_name)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (doc_id, title, text, str(path), metadata["doc_type"], metadata["project_name"], metadata["topic_name"]),
    )
    return {
        "path": str(path),
        "doc_type": metadata["doc_type"],
        "source_group": source_group,
        "project_name": metadata["project_name"],
        "topic_name": metadata["topic_name"],
        "extract_status": status,
    }


def remove_document(conn: sqlite3.Connection, path: str) -> None:
    row = conn.execute("SELECT id FROM docs WHERE path = ?", (path,)).fetchone()
    if not row:
        return
    doc_id = int(row["id"])
    conn.execute("DELETE FROM docs_fts WHERE rowid = ?", (doc_id,))
    conn.execute("DELETE FROM docs WHERE id = ?", (doc_id,))


def build_index() -> dict[str, Any]:
    db = index_db_path()
    if db.exists():
        db.unlink()
    conn = db_connect()
    indexed = 0
    extracted = {"ok": 0, "truncated": 0, "no_text": 0, "unsupported": 0, "error": 0}
    for path in sorted(iter_candidate_files()):
        record = upsert_document(conn, path)
        indexed += 1
        extracted[record["extract_status"]] = extracted.get(record["extract_status"], 0) + 1
    conn.commit()
    conn.close()
    state = {
        "last_build_at": iso_now(),
        "last_sync_at": iso_now(),
        "doc_count": indexed,
        "dirty_count": 0,
        "extract_status_counts": extracted,
    }
    save_state(state)
    return state


def sync_index() -> dict[str, Any]:
    conn = db_connect()
    current_files = {str(path): path for path in sorted(iter_candidate_files())}
    rows = conn.execute(
        "SELECT path, mtime, size, content_hash, doc_type, source_group, heading, line_start, line_end FROM docs"
    ).fetchall()
    existing = {str(row["path"]): row for row in rows}
    changed = 0
    removed = 0
    for path_str, row in existing.items():
        if path_str not in current_files:
            remove_document(conn, path_str)
            removed += 1
    for path_str, path in current_files.items():
        stat = path.stat()
        row = existing.get(path_str)
        if row and float(row["mtime"]) == stat.st_mtime and int(row["size"]) == stat.st_size:
            current_metadata = classify_document(path)
            expected_group = classify_source_group(current_metadata["doc_type"], path)
            if (
                str(row["doc_type"]) == current_metadata["doc_type"]
                and str(row["source_group"]) == expected_group
                and str(row["heading"]).strip()
                and int(row["line_end"] or 0) >= int(row["line_start"] or 0) >= 0
                and not (int(row["line_start"] or 0) == 0 and int(row["line_end"] or 0) == 0)
            ):
                continue
        upsert_document(conn, path)
        changed += 1
    conn.commit()
    doc_count = int(conn.execute("SELECT COUNT(*) FROM docs").fetchone()[0])
    conn.close()
    state = load_state()
    state.update(
        {
            "last_sync_at": iso_now(),
            "doc_count": doc_count,
            "dirty_count": 0,
            "changed_count": changed,
            "removed_count": removed,
        }
    )
    save_state(state)
    return state


def search_index(
    query: str,
    *,
    doc_type: str = "",
    project_name: str = "",
    topic_name: str = "",
    hotset_paths: Iterable[str] | None = None,
    limit: int = 10,
) -> list[dict[str, Any]]:
    if not query.strip():
        return []
    conn = db_connect()
    group_weight_sql = "CASE docs.source_group " + " ".join(
        f"WHEN '{group_name}' THEN {weight}" for group_name, weight in SOURCE_GROUP_WEIGHT.items()
    ) + " ELSE 99 END"
    doc_weight_sql = "CASE docs.doc_type " + " ".join(
        f"WHEN '{doc_type_name}' THEN {weight}" for doc_type_name, weight in DOC_TYPE_WEIGHT.items()
    ) + " ELSE 99 END"
    filters = []
    params: list[Any] = [query]
    if doc_type:
        filters.append("docs.doc_type = ?")
        params.append(doc_type)
    if project_name:
        filters.append("docs.project_name = ?")
        params.append(project_name)
    if topic_name:
        filters.append("docs.topic_name = ?")
        params.append(topic_name)
    where = " AND ".join(filters)
    sql = f"""
        SELECT
            docs.path,
            docs.title,
            docs.doc_type,
            docs.source_group,
            docs.heading,
            docs.line_start,
            docs.line_end,
            docs.project_name,
            docs.topic_name,
            docs.extract_status,
            docs.content,
            snippet(docs_fts, 1, '[', ']', '…', 16) AS excerpt,
            bm25(docs_fts) AS score
        FROM docs_fts
        JOIN docs ON docs.id = docs_fts.rowid
        WHERE docs_fts MATCH ?
        {'AND ' + where if where else ''}
        ORDER BY {group_weight_sql} ASC, {doc_weight_sql} ASC, bm25(docs_fts) ASC
        LIMIT ?
    """
    fetch_limit = limit * 5 if hotset_paths else limit
    params.append(fetch_limit)
    try:
        rows = conn.execute(sql, params).fetchall()
    except sqlite3.OperationalError:
        rows = conn.execute(sql, [f'"{query}"', *params[1:]]).fetchall()
    finally:
        conn.close()
    results: list[dict[str, Any]] = []
    for row in rows:
        hotset = is_hotset_path(row["path"], hotset_paths)
        heading, line_start, line_end, excerpt = select_best_chunk(
            content=str(row["content"] or ""),
            title=str(row["title"] or ""),
            fallback_heading=str(row["heading"] or row["title"] or ""),
            fallback_line_start=int(row["line_start"] or 0),
            fallback_line_end=int(row["line_end"] or 0),
            base_excerpt=str(row["excerpt"] or ""),
            query=query,
        )
        results.append(
            {
                "path": row["path"],
                "title": row["title"],
                "doc_type": row["doc_type"],
                "source_group": row["source_group"],
                "heading": heading,
                "line_start": line_start,
                "line_end": line_end,
                "project_name": row["project_name"],
                "topic_name": row["topic_name"],
                "extract_status": row["extract_status"],
                "score": row["score"],
                "excerpt": excerpt,
                "is_hotset": hotset,
                "pin_reason": "hotset_path" if hotset else "",
            }
        )
    if hotset_paths:
        results.sort(
            key=lambda item: (
                0 if item["is_hotset"] else 1,
                SOURCE_GROUP_WEIGHT.get(str(item.get("source_group", "")), 99),
                DOC_TYPE_WEIGHT.get(str(item.get("doc_type", "")), 99),
                float(item.get("score", 0)),
            )
        )
    return results[:limit]


def get_document(path: str) -> dict[str, Any]:
    conn = db_connect()
    row = conn.execute(
        """
        SELECT path, title, doc_type, project_name, topic_name, mtime, size, content_hash, extract_status, content, indexed_at
        , source_group, heading, line_start, line_end
        FROM docs WHERE path = ?
        """,
        (path,),
    ).fetchone()
    conn.close()
    if not row:
        raise FileNotFoundError(path)
    return dict(row)


def status() -> dict[str, Any]:
    state = load_state()
    db_exists = index_db_path().exists()
    doc_count = 0
    if db_exists:
        conn = db_connect()
        doc_count = int(conn.execute("SELECT COUNT(*) FROM docs").fetchone()[0])
        conn.close()
    return {
        "db_path": str(index_db_path()),
        "db_exists": db_exists,
        "last_build_at": state.get("last_build_at"),
        "last_sync_at": state.get("last_sync_at"),
        "doc_count": doc_count,
        "dirty_count": state.get("dirty_count", 0),
    }


def cmd_build_index(_args: argparse.Namespace) -> int:
    print(json.dumps(build_index(), ensure_ascii=False))
    return 0


def cmd_sync_index(_args: argparse.Namespace) -> int:
    print(json.dumps(sync_index(), ensure_ascii=False))
    return 0


def cmd_search(args: argparse.Namespace) -> int:
    print(
        json.dumps(
            search_index(
                args.query,
                doc_type=args.doc_type,
                project_name=args.project_name,
                topic_name=args.topic_name,
                limit=args.limit,
            ),
            ensure_ascii=False,
        )
    )
    return 0


def cmd_get(args: argparse.Namespace) -> int:
    try:
        payload = get_document(args.path)
    except FileNotFoundError:
        print(json.dumps({"error": f"Document not indexed: {args.path}"}, ensure_ascii=False))
        return 1
    print(json.dumps(payload, ensure_ascii=False))
    return 0


def cmd_status(_args: argparse.Namespace) -> int:
    print(json.dumps(status(), ensure_ascii=False))
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="workspace-hub local retrieval index")
    subparsers = parser.add_subparsers(dest="command", required=True)

    build = subparsers.add_parser("build-index")
    build.set_defaults(func=cmd_build_index)

    sync = subparsers.add_parser("sync-index")
    sync.set_defaults(func=cmd_sync_index)

    search = subparsers.add_parser("search")
    search.add_argument("--query", required=True)
    search.add_argument("--doc-type", default="")
    search.add_argument("--project-name", default="")
    search.add_argument("--topic-name", default="")
    search.add_argument("--limit", type=int, default=10)
    search.set_defaults(func=cmd_search)

    get_cmd = subparsers.add_parser("get")
    get_cmd.add_argument("--path", required=True)
    get_cmd.set_defaults(func=cmd_get)

    status_cmd = subparsers.add_parser("status")
    status_cmd.set_defaults(func=cmd_status)
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    return args.func(args)


if __name__ == "__main__":
    raise SystemExit(main())
