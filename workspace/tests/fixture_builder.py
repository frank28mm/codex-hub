from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from openpyxl import Workbook


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def write_sample_feishu_resources(control_root: Path) -> None:
    write_text(
        control_root / "feishu_resources.yaml",
        (
            "version: 1\n"
            "defaults:\n"
            "  owner_open_id: \"ou_owner\"\n"
            "  calendar_id: \"cal_default\"\n"
            "  personal_calendar_id: \"\"\n"
            "  calendar_create_default_route: \"invite_meeting\"\n"
            "  personal_reminder_target: \"task\"\n"
            "  doc_folder_token: \"fld_default\"\n"
            "  oauth_scopes:\n"
            "    - \"bitable:app\"\n"
            "    - \"base:record:create\"\n"
            "    - \"base:record:update\"\n"
            "    - \"base:record:delete\"\n"
            "    - \"contact:contact.base:readonly\"\n"
            "    - \"task:task:write\"\n"
            "    - \"task:tasklist:read\"\n"
            "    - \"task:tasklist:write\"\n"
            "  meeting:\n"
            "    calendar_id: \"cal_meeting\"\n"
            "    timezone: \"Asia/Shanghai\"\n"
            "    duration_minutes: 30\n"
            "    attendee_ability: \"can_modify_event\"\n"
            "    visibility: \"default\"\n"
            "aliases:\n"
            "  chats:\n"
            "    产品群: \"oc_group_123\"\n"
            "  users:\n"
            "    Operator:\n"
            "      open_id: \"ou_owner\"\n"
            "      email: \"operator@example.com\"\n"
            "  calendars:\n"
            "    默认: \"cal_default\"\n"
            "  doc_folders:\n"
            "    报告: \"fld_reports\"\n"
            "  tables:\n"
            "    书单:\n"
            "      app_token: \"app_book\"\n"
            "      table_id: \"tbl_book\"\n"
            "  tasklists: {}\n"
            "projection:\n"
            "  app:\n"
            "    alias: \"codex_hub_projection\"\n"
            "    name: \"Codex Hub 项目任务看板\"\n"
            "    app_token: \"\"\n"
            "    folder_token: \"\"\n"
            "  tables:\n"
            "    projects_overview:\n"
            "      alias: \"codex_hub_projects_overview\"\n"
            "      name: \"项目总览\"\n"
            "      table_id: \"\"\n"
            "      default_view_name: \"全部项目\"\n"
            "    tasks_current:\n"
            "      alias: \"codex_hub_tasks_current\"\n"
            "      name: \"当前任务\"\n"
            "      table_id: \"\"\n"
            "      default_view_name: \"全部任务\"\n"
            "  views:\n"
            "    projects_overview:\n"
            "      - name: \"全部项目\"\n"
            "        type: \"grid\"\n"
            "      - name: \"按状态看板\"\n"
            "        type: \"kanban\"\n"
            "      - name: \"按优先级\"\n"
            "        type: \"grid\"\n"
            "      - name: \"最近更新\"\n"
            "        type: \"grid\"\n"
            "      - name: \"需关注项目\"\n"
            "        type: \"grid\"\n"
            "    tasks_current:\n"
            "      - name: \"全部任务\"\n"
            "        type: \"grid\"\n"
            "      - name: \"按状态看板\"\n"
            "        type: \"kanban\"\n"
            "      - name: \"按项目分组\"\n"
            "        type: \"kanban\"\n"
            "      - name: \"阻塞项\"\n"
            "        type: \"grid\"\n"
            "      - name: \"最近更新任务\"\n"
            "        type: \"grid\"\n"
        ),
    )


def escape_pdf_text(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def write_minimal_pdf(path: Path, text: str) -> None:
    stream = f"BT /F1 18 Tf 72 720 Td ({escape_pdf_text(text)}) Tj ET".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    parts = [b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"]
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(sum(len(part) for part in parts))
        parts.append(f"{index} 0 obj\n".encode("ascii"))
        parts.append(obj)
        parts.append(b"\nendobj\n")
    xref_offset = sum(len(part) for part in parts)
    parts.append(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    parts.append(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        parts.append(f"{offset:010d} 00000 n \n".encode("ascii"))
    parts.append(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(b"".join(parts))


def build_sample_environment(root: Path) -> dict[str, Path]:
    workspace_root = root / "workspace-hub-fixture"
    vault_root = root / "vault"
    projects_root = workspace_root / "projects"
    reports_root = workspace_root / "reports"
    runtime_root = workspace_root / "runtime"
    control_root = workspace_root / "control"

    (vault_root / "01_working").mkdir(parents=True, exist_ok=True)
    (vault_root / "02_episodic" / "daily").mkdir(parents=True, exist_ok=True)
    (vault_root / "03_semantic" / "material_routes").mkdir(parents=True, exist_ok=True)
    (vault_root / "03_semantic" / "systems").mkdir(parents=True, exist_ok=True)
    (vault_root / "03_semantic" / "projects").mkdir(parents=True, exist_ok=True)
    (vault_root / "07_dashboards").mkdir(parents=True, exist_ok=True)
    (vault_root / "07_dashboards" / "materials").mkdir(parents=True, exist_ok=True)
    reports_root.mkdir(parents=True, exist_ok=True)
    projects_root.mkdir(parents=True, exist_ok=True)
    runtime_root.mkdir(parents=True, exist_ok=True)
    control_root.mkdir(parents=True, exist_ok=True)
    write_sample_feishu_resources(control_root)

    write_text(
        vault_root / "PROJECT_REGISTRY.md",
        "# PROJECT_REGISTRY\n\n## Registry Data\n\n<!-- PROJECT_REGISTRY_DATA_START -->\n```json\n[\n  {\n    \"project_name\": \"SampleProj\",\n    \"aliases\": [],\n    \"path\": \"/tmp/SampleProj\",\n    \"status\": \"active\",\n    \"summary_note\": \"sample\"\n  }\n]\n```\n<!-- PROJECT_REGISTRY_DATA_END -->\n",
    )
    write_text(vault_root / "ACTIVE_PROJECTS.md", "# ACTIVE_PROJECTS\n")
    write_text(vault_root / "NEXT_ACTIONS.md", "# NEXT_ACTIONS\n")
    write_text(
        vault_root / "01_working" / "SampleProj-项目板.md",
        "---\nboard_type: project\nproject_name: SampleProj\nstatus: active\npriority: high\nupdated_at: 2026-03-11\npurpose: sample project board\n---\n\n# Sample Board\n\nAlpha Board task line\n",
    )
    write_text(
        vault_root / "01_working" / "SampleProj-需求-跟进板.md",
        "---\nboard_type: topic\nproject_name: SampleProj\ntopic_name: 需求\ntopic_key: demand\nrollup_target: /tmp/SampleProj-board.md\nupdated_at: 2026-03-11\npurpose: sample topic board\n---\n\n# Sample Topic\n\nTopic Retrieval Marker\n",
    )
    write_text(
        vault_root / "03_semantic" / "systems" / "workspace-hub.md",
        "# workspace-hub\n\nControl System Rule\n",
    )
    write_text(
        vault_root / "03_semantic" / "projects" / "SampleProj.md",
        "---\nproject_name: SampleProj\nstatus: active\npriority: high\npath: /tmp/SampleProj\nupdated_at: 2026-03-11\nsummary: Sample project summary.\n---\n\n# SampleProj\n\nProject Summary Memory\n",
    )
    write_text(
        vault_root / "03_semantic" / "projects" / "SampleProj-机制建议.md",
        "# SampleProj 建议\n\nAdvisory note\n",
    )
    write_text(
        vault_root / "02_episodic" / "daily" / "2026-03-10.md",
        "# 2026-03-10\n\nDaily note with repeated keyword ledger\n",
    )
    write_text(vault_root / "07_dashboards" / "PROJECTS.md", "# PROJECTS\n")
    write_text(reports_root / "system-overview.md", "# Report\n\nSystem Report Marker\n")

    sample_project = projects_root / "SampleProj"
    sample_project.mkdir(parents=True, exist_ok=True)
    write_text(sample_project / "guide.md", "# Guide\n\nProject Document Marker\n")
    write_text(sample_project / "notes.txt", "Plain text note for retrieval.\n")
    write_text(sample_project / "src" / "app.py", "SECRET_CODE_ONLY = 'should not be indexed'\n")

    write_minimal_pdf(sample_project / "slides.pdf", "PDF Fixture Marker")

    doc = Document()
    doc.add_heading("Docx Title", level=1)
    doc.add_paragraph("Docx Fixture Marker")
    doc.save(sample_project / "guide.docx")

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "SheetOne"
    sheet["A1"] = "Budget"
    sheet["B1"] = "Approved"
    sheet["A2"] = "XLSX Fixture Marker"
    workbook.save(sample_project / "budget.xlsx")

    with (sample_project / "items.csv").open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(["name", "value"])
        writer.writerow(["CSV Fixture Marker", "42"])

    write_text(
        vault_root / "03_semantic" / "material_routes" / "SampleProj.md",
        (
            "# SampleProj 材料路由\n\n"
            "这是 SampleProj 的材料路由配置。\n\n"
            "<!-- MATERIAL_ROUTE_CONFIG_START -->\n"
            "```json\n"
            "{\n"
            f"  \"project_material_roots\": [\"{sample_project.as_posix()}\"],\n"
            f"  \"report_roots\": [\"{reports_root.as_posix()}\"],\n"
            f"  \"deliverable_roots\": [\"{(sample_project / 'deliverables').as_posix()}\"],\n"
            f"  \"hotset_paths\": [\"{(sample_project / 'guide.md').as_posix()}\"] ,\n"
            f"  \"ignore_paths\": [\"{(sample_project / 'notes.txt').as_posix()}\"] ,\n"
            "  \"preferred_queries\": [\"Project Document Marker\", \"System Report Marker\"],\n"
            f"  \"allow_roots\": [\"{sample_project.as_posix()}\", \"{reports_root.as_posix()}\"]\n"
            "}\n"
            "```\n"
            "<!-- MATERIAL_ROUTE_CONFIG_END -->\n"
        ),
    )

    return {
        "workspace_root": workspace_root,
        "vault_root": vault_root,
        "projects_root": projects_root,
        "reports_root": reports_root,
        "runtime_root": runtime_root,
        "control_root": control_root,
        "sample_project": sample_project,
    }
